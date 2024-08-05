
class Generator():
    def Form1(self, n):
        import docx
        import random
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        name = open('dataFile_security/Data_FORM1/이름.txt', 'r', encoding = 'UTF8')
        name_data = name.read()
        name = name_data.split('\n')

        company = open('dataFile_security/Data_FORM1/회사명.txt', 'r', encoding = 'UTF8')
        company_data = company.read()
        company = company_data.split('\n')

        address = open('dataFile_security/Data_FORM1/주소.txt', 'r', encoding = 'UTF8')
        address_data = address.read()
        address = address_data.split('\n')

        belong = open('dataFile_security/Data_FORM1/소속(부서).txt', 'r', encoding = 'UTF8')
        belong_data = belong.read()
        belong = belong_data.split('\n')

        address2 = open('dataFile_security/Data_FORM1/주소.txt', 'r', encoding = 'UTF8')
        address_data2 = address2.read()
        address2 = address_data2.split('\n')

        birth_num = open('dataFile_security/Data_FORM1/주민번호.txt', 'r', encoding = 'UTF8')
        birth_num_data = birth_num.read()
        birth_num = birth_num_data.split('\n')

        date = open('dataFile_security/Data_FORM1/날짜.txt', 'r', encoding = 'UTF8')
        date_data = date.read()
        date = date_data.split('\n')

        for i in range(n):
            Form1 = docx.Document('formFile_security/Form1.docx')

            company_index = company[random.randrange(len(company))]
            belong_index = belong[random.randrange(len(belong))]
            date_index = date[random.randrange(len(date))]
            address_index = address[random.randrange(len(address))]
            address2_index = address2[random.randrange(len(address2))]
            name_index = name[random.randrange(len(name))]
            birth_num_index = birth_num[random.randrange(len(birth_num))]


            Form1.add_paragraph(' 본인 %s 은 %s 에 소재하고 있는 %s 에서 퇴직함에 있어 다음 사항을 숙지하고 이를 이행하지 않을 경우 관계 법령에 의거 처벌받을 것은 물론 %s 에 손해를 입힐 경우에는 그 손해액을 변상할 것을 엄숙히 서약합니다.' %(name_index, address_index, company_index, company_index), style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('1. %s 에 근무 중 지득한 국가보안 등에 관한 제반 비밀과 직무상 지득한 과학기술정보 관련 제반 비밀사항 및 주요 기술비밀을 관련 법령, 인사규정 제 10조, 취업규칙 제 2조의 규정에 따라 일체 누설하거나 도용하지 않는다.' %company_index, style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('2. %s 에 근무 중의 모든 발명, 고안, 창작 및 발견 등에 대하여 %s %s 에게 이를 공개, 양도할 것에 동의하고 그 절차에 적극 협력한다.' %(company_index, company_index, belong_index), style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('3. %s 에 근무 중의 모든 연구자료 및 연구결과 보고서, 설계서, 청사진 등과 보조기억장치 등에 대하여는 누락없이 %s %s에게 인계하고 이를 소지하거나 유출하지 않는다.' %(company_index, company_index, belong_index), style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('4. 퇴직 후 2년 간은 %s 의 사전 승인 없이 %s 의 연구자료, 연구결과 등과 직무 발명, 고안, 창작 및 발견사항 등의 지적재산권을 이용하여 자신 또는 제 3자를 위하여 창업하거나, 기업체에 전직, 동업 또는 자문하지 않는다.' %(company_index, company_index), style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('5. 위 사항을 위반하는 경우에는 관련 법규(국가보안법, 형법, 부정경쟁방지 및 영업비밀보호에 관련 법률)에 따른 어떠한 처벌도 감수한다.', style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('%s' %date_index, style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('서약인 주소 : %s' %address2_index, style='글씨9')
            Form1.add_paragraph('서약인 주민등록번호 : %s' %birth_num_index, style='글씨9')
            Form1.add_paragraph('서약인 성명 : %s' %name_index, style='글씨9')
            Form1.add_paragraph('', style='글씨9')
            Form1.add_paragraph('%s 귀하' %company_index, style='볼드')

            Form1.paragraphs[14].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Form1.paragraphs[20].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            Form1.save('../../_inputData/raw/documents/security_form1_{}.docx'.format(str(i+1)))

    def Form2(self, n):
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import random

        name = open('dataFile_security/Data_FORM2/이름.txt', 'r', encoding = 'UTF8')
        name_data = name.read()
        name = name_data.split('\n')

        company1 = open('dataFile_security/Data_FORM2/회사명.txt', 'r', encoding = 'UTF8')
        company1_data = company1.read()
        company1 = company1_data.split('\n')

        company2 = open('dataFile_security/Data_FORM2/회사명.txt', 'r', encoding = 'UTF8')
        company2_data = company2.read()
        company2 = company2_data.split('\n')

        address = open('dataFile_security/Data_FORM2/주소.txt', 'r', encoding = 'UTF8')
        address_data = address.read()
        address = address_data.split('\n')

        belong = open('dataFile_security/Data_FORM2/소속(부서).txt', 'r', encoding = 'UTF8')
        belong_data = belong.read()
        belong = belong_data.split('\n')

        period = open('dataFile_security/Data_FORM2/참여기간(기간1).txt', 'r', encoding = 'UTF8')
        period_data2 = period.read()
        period = period_data2.split('\n')

        date = open('dataFile_security/Data_FORM2/날짜.txt', 'r', encoding = 'UTF8')
        date_data = date.read()
        date = date_data.split('\n')

        status = open('dataFile_security/Data_FORM2/직위.txt', 'r', encoding = 'UTF8')
        status_data = status.read()
        status = status_data.split('\n')

        for i in range(n):
            Form2 = docx.Document('formFile_security/Form2.docx')
            company1_index = company1[random.randrange(len(company1))]
            company2_index = company2[random.randrange(len(company2))]
            belong_index = belong[random.randrange(len(belong))]
            date_index = date[random.randrange(len(date))]
            address_index = address[random.randrange(len(address))]
            period_index = period[random.randrange(len(period))]
            name_index = name[random.randrange(len(name))]

            if company1_index != company2_index:
                Form2.add_paragraph(' %s %s 고객님 귀중' %(company2_index, belong_index))
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph(' %s 에 위치한 본사 %s 은 %s %s (이하 %s 이라 한다) 고객으로부터 %s 견적 의뢰 받은 개인민감정보 데이터 분석 서버 관련 문서에 대한 보안 유지를 아래와 같이 서약합니다.' %(address_index, company1_index, company2_index, belong_index, company2_index, date_index))
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph('1. 의뢰 문서 내용에 대한 보안을 유지하는 것을 기본적으로 %s 견적 발송 이후 모든 해당 자료는 %s 의 동의 없이는 외부로 유출하지 않을 것을 확인 드립니다.' %(date_index, company2_index))
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph('2. 원문 자료 또는 번역문은 번역 외 기타 목적으로 사용하지 않으며, 필요 시 %s의 동의 하에 사용될 것을 확인 드립니다.' %company2_index)
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph('3. 번역 완료 후 디지털 파일 또는 문서 사본(영문 및 한글 번역본 모두 포함)은 A/S 유지를 위해 %s 기간 동안 보관하며, 기한 후 모든 자료는 복구 불능한 상태로 삭제 또는 파기할 것을 확인 드립니다. (필요시 번역 완료 후 즉시 삭제 및 파기 가능)' %period_index)
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph('4. %s 의 실책으로 비밀 준수 사항을 위반하여 %s에 손해를 미친 경우 모든 손해 사항에 대하여 배상할 것을 확인 드립니다.' %(company1_index, company2_index))
                Form2.add_paragraph('', style='글씨5')
                Form2.add_paragraph('5. 본 서약서는 %s 및 %s 의 고객 체결 완료일부터 효력이 발생합니다.' %(company2_index, company1_index))
                Form2.add_paragraph('', style='글씨7')
                Form2.add_paragraph('%s' %date_index)
                Form2.add_paragraph('')
                Form2.add_paragraph('고객 대표자 %s %s %s' %(company2_index, belong_index, name_index))
                Form2.add_paragraph('%s 제공' %company1_index)

                Form2.paragraphs[16].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                Form2.paragraphs[18].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                Form2.paragraphs[19].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                Form2.save('../../_inputData/raw/documents/security_form2_{}.docx'.format(str(i+1)))

    def Form3(self, n):
        import docx
        import random
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        name = open('dataFile_security/Data_FORM3/이름.txt', 'r', encoding = 'UTF8')
        name_data = name.read()
        name = name_data.split('\n')

        company = open('dataFile_security/Data_FORM3/회사명.txt', 'r', encoding = 'UTF8')
        company_data = company.read()
        company = company_data.split('\n')

        address = open('dataFile_security/Data_FORM3/주소.txt', 'r', encoding = 'UTF8')
        address_data = address.read()
        address = address_data.split('\n')

        belong = open('dataFile_security/Data_FORM3/소속(부서).txt', 'r', encoding = 'UTF8')
        belong_data = belong.read()
        belong = belong_data.split('\n')

        date = open('dataFile_security/Data_FORM3/날짜.txt', 'r', encoding = 'UTF8')
        date_data = date.read()
        date = date_data.split('\n')

        status = open('dataFile_security/Data_FORM3/직위.txt', 'r', encoding = 'UTF8')
        status_data = status.read()
        status = status_data.split('\n')

        for i in range(n):
            Form3 = docx.Document('formFile_security/Form3.docx')

            company_index = company[random.randrange(len(company))]
            belong_index = belong[random.randrange(len(belong))]
            date_index = date[random.randrange(len(date))]
            address_index = address[random.randrange(len(address))]
            name_index = name[random.randrange(len(name))]
            status_index = status[random.randrange(len(status))]

            Form3.add_paragraph('소속 : %s %s' %(company_index, belong_index))
            Form3.add_paragraph('직위 : %s' %status_index)
            Form3.add_paragraph('성명 : %s' %name_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph(' 본인 %s 은 %s 소재지에 있는 %s 의 영업비밀 관리규정을 충분히 숙지, 이해하였으며 다음의 사항을 준수할 것을 엄숙히 서약합니다.' %(name_index, address_index, company_index))
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('1. %s 의 영업비밀 관리 규정과 이에 관련한 명령을 성실히 이행하겠습니다.' %company_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('2. %s 의 영업비밀은 재직 중은 물론 퇴직 후에도 회사의 허가 없이 사용하거나 제 3자에게 무단 누설하거나 경쟁회사에 유출하지 않겠습니다.' %company_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('3. 본인 %s 이 알고 있는 제 3자의 영업비밀은 여하한 일이 있어도 비밀 보유자의 승낙 없이 회사에 공개하거나 회사의 업무에 부정하게 사용하지 않겠습니다.' %name_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('4. %s 재직 시 지득한 영업비밀과 관련하여 경쟁회사에서는 이와 동일한 %s 업무를 담당하지 않겠습니다.' %(company_index, belong_index))
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('5. 재직 시는 물론 퇴직 후에도 %s 재직 시 지득한 영업비밀을 가지고 창업을 하거나 경쟁 회사에 전직 또는 동업을 하지 않겠습니다.' %company_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('6. 만약 이 서약서에 위반할 경우에는 부정경쟁방지법의 관련 규정과 %s 의 영업비밀관리 규정에 의한 어떠한 조치도 감수하겠습니다.' %company_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('%s' %date_index)
            Form3.add_paragraph('서약인 : %s (인)' %name_index)
            Form3.add_paragraph('', style='글씨5')
            Form3.add_paragraph('%s 귀하' %company_index, style='볼드')

            Form3.paragraphs[20].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Form3.paragraphs[21].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            Form3.paragraphs[23].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            Form3.save('../../_inputData/raw/documents/security_form3_{}.docx'.format(str(i+1)))


Generator().Form1(10)

Generator().Form2(10)

# Generator().Form3()




