def ldcc2_st(x):
    from docx import Document
    import kss
    first_data = []
    second_data = []
    third_data = []
    final_result = []
    document = Document(x)

    tables = document.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    first_data.append(paragraph.text)

    for i in range(0, len(first_data)):
        first_data[i] = first_data[i].replace('\n', ' ')
        first_data[i] = first_data[i].replace('  ', ' ')
        first_data[i] = first_data[i].replace('- 다음 -', '')
    for i in range(5, len(first_data)):
        for sent in kss.split_sentences(first_data[i]):
            second_data.append(sent)

    for i in range(0, len(second_data)-11):
        final_result.append(second_data[i])

    for i in range(3, 6):
        third_data.append(second_data[i])
    final_result.append(' '.join(third_data))

    for i in range(6, 11):
        final_result.append(second_data[i])
    third_data = []
    for i in range(11, len(second_data)):
        third_data.append(second_data[i])
    final_result.append(' '.join(third_data))

    return final_result


