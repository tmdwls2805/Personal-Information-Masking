#def ldcc2_st(x):
from docx import Document
first_data = []
second_data = []
third_data = []
final_result = []
document = Document('../../_inputData/raw/documents/ldcc_form2_01.docx')

tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                first_data.append(paragraph.text)
print(first_data)
for i in range(0, len(first_data)):
    if first_data[i] != '':
        second_data.append(first_data[i])

first_data = []
for i in range(13, len(second_data)):
    first_data.append(second_data[i])
second_data = []
for i in range(0, len(first_data)):
    if first_data[i] != '- 다 음 -':
        second_data.append(first_data[i])
first_data = []
for i in range(0, len(second_data)-7):
    first_data.append(second_data[i])
for i in range(len(second_data)-7, len(second_data)):
    third_data.append(second_data[i])
second_data = []
second_data = ' '.join(third_data)

for i in range(0, len(first_data)):
    final_result.append(first_data[i])
final_result.append(second_data)
print(final_result)
#return final_result