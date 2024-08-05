def ldcc2(x):
    from _utils.common.Mecab import Mecab
    import kss
    from docx import Document
    first_data = []
    second_data = []
    third_data = []
    final_result = []
    document = Document(x)

    tables = document.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    first_data.append(paragraph.text)

    for i in range(0, len(first_data)):
        first_data[i] = first_data[i].replace('\n', ' ')
        first_data[i] = first_data[i].replace('  ', ' ')
        first_data[i] = first_data[i].replace('- 다음 -', '')
    for i in range(5, len(first_data)):
        for sent in kss.split_sentences(first_data[i]):
            second_data.append(sent)

    for i in range(0, len(second_data)-11):
        final_result.append(second_data[i])

    for i in range(3, 6):
        third_data.append(second_data[i])
    final_result.append(' '.join(third_data))

    for i in range(6, 11):
        final_result.append(second_data[i])
    third_data = []
    for i in range(11, len(second_data)):
        third_data.append(second_data[i])
    final_result.append(' '.join(third_data))

    me = Mecab()
    inter_result = []
    word_list = []
    pos_list = []


    for word in final_result:
        word_set = me.pos(word)
        inter_result.append(word_set)

    for i in range(0, len(inter_result)):
        pos_result = []
        word_result = []
        for j in range(0, len(inter_result[i])):
            word_result.append(inter_result[i][j][0])
            pos_result.append(inter_result[i][j][1])
        word_list.append(word_result)
        pos_list.append(pos_result)

    final_result = []
    for i in range(0, len(word_list)):
        word_pos = []
        for j in range(0, len(word_list[i])):
            word_pos.append(word_list[i][j] + " " + pos_list[i][j])
        final_result.append(word_pos)

    return final_result